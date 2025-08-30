import json
import csv
import xml.etree.ElementTree as ET
from typing import List, Dict, Any, Optional
from datetime import datetime
import logging
from pathlib import Path
import zipfile
import io
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class ExportService:
    """导出服务"""
    
    def __init__(self):
        self.supported_formats = ["json", "csv", "xml", "pdf", "docx", "zip"]
        self.styles = getSampleStyleSheet()
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的导出格式"""
        return self.supported_formats
    
    def export_documents(self, documents: List[Dict], format: str, **kwargs) -> bytes:
        """导出文档列表"""
        try:
            if format == "json":
                return self._export_json(documents, **kwargs)
            elif format == "csv":
                return self._export_csv(documents, **kwargs)
            elif format == "xml":
                return self._export_xml(documents, **kwargs)
            elif format == "pdf":
                return self._export_pdf(documents, **kwargs)
            elif format == "docx":
                return self._export_docx(documents, **kwargs)
            elif format == "zip":
                return self._export_zip(documents, **kwargs)
            else:
                raise ValueError(f"不支持的导出格式: {format}")
        except Exception as e:
            logger.error(f"导出失败: {e}")
            raise
    
    def export_lesson_plans(self, lesson_plans: List[Dict], format: str, **kwargs) -> bytes:
        """导出教案列表"""
        try:
            if format == "json":
                return self._export_json(lesson_plans, **kwargs)
            elif format == "csv":
                return self._export_csv(lesson_plans, **kwargs)
            elif format == "xml":
                return self._export_xml(lesson_plans, **kwargs)
            elif format == "pdf":
                return self._export_pdf(lesson_plans, **kwargs)
            elif format == "docx":
                return self._export_docx(lesson_plans, **kwargs)
            elif format == "zip":
                return self._export_zip(lesson_plans, **kwargs)
            else:
                raise ValueError(f"不支持的导出格式: {format}")
        except Exception as e:
            logger.error(f"导出失败: {e}")
            raise
    
    def _export_json(self, data: List[Dict], **kwargs) -> bytes:
        """导出为JSON格式"""
        export_data = {
            "export_time": datetime.now().isoformat(),
            "total_count": len(data),
            "data": data
        }
        
        # 处理特殊字段
        for item in export_data["data"]:
            if "created_at" in item and isinstance(item["created_at"], datetime):
                item["created_at"] = item["created_at"].isoformat()
            if "updated_at" in item and isinstance(item["updated_at"], datetime):
                item["updated_at"] = item["updated_at"].isoformat()
        
        return json.dumps(export_data, ensure_ascii=False, indent=2, default=str).encode('utf-8')
    
    def _export_csv(self, data: List[Dict], **kwargs) -> bytes:
        """导出为CSV格式"""
        if not data:
            return b""
        
        # 获取所有字段
        fields = set()
        for item in data:
            fields.update(item.keys())
        fields = sorted(list(fields))
        
        # 创建CSV内容
        output = io.StringIO()
        writer = csv.DictWriter(output, fieldnames=fields)
        
        # 写入表头
        writer.writeheader()
        
        # 写入数据
        for item in data:
            # 处理特殊字段
            row = {}
            for field in fields:
                value = item.get(field, "")
                if isinstance(value, (list, tuple)):
                    value = ", ".join(str(v) for v in value)
                elif isinstance(value, datetime):
                    value = value.isoformat()
                elif value is None:
                    value = ""
                row[field] = str(value)
            writer.writerow(row)
        
        return output.getvalue().encode('utf-8')
    
    def _export_xml(self, data: List[Dict], **kwargs) -> bytes:
        """导出为XML格式"""
        root = ET.Element("export")
        root.set("export_time", datetime.now().isoformat())
        root.set("total_count", str(len(data)))
        
        for item in data:
            item_elem = ET.SubElement(root, "item")
            for key, value in item.items():
                if isinstance(value, (list, tuple)):
                    # 处理列表字段
                    list_elem = ET.SubElement(item_elem, key)
                    for v in value:
                        sub_elem = ET.SubElement(list_elem, "value")
                        sub_elem.text = str(v)
                elif isinstance(value, dict):
                    # 处理字典字段
                    dict_elem = ET.SubElement(item_elem, key)
                    for k, v in value.items():
                        sub_elem = ET.SubElement(dict_elem, k)
                        sub_elem.text = str(v)
                else:
                    # 处理普通字段
                    elem = ET.SubElement(item_elem, key)
                    if value is not None:
                        elem.text = str(value)
        
        return ET.tostring(root, encoding='utf-8', xml_declaration=True)
    
    def _export_pdf(self, data: List[Dict], **kwargs) -> bytes:
        """导出为PDF格式"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # 标题
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # 居中
        )
        title = Paragraph("数据导出报告", title_style)
        story.append(title)
        
        # 导出信息
        info_style = ParagraphStyle(
            'Info',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=20
        )
        info_text = f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>"
        info_text += f"总数量: {len(data)}<br/>"
        info = Paragraph(info_text, info_style)
        story.append(info)
        
        if data:
            # 创建表格
            headers = list(data[0].keys())
            table_data = [headers]  # 表头
            
            for item in data:
                row = []
                for header in headers:
                    value = item.get(header, "")
                    if isinstance(value, (list, tuple)):
                        value = ", ".join(str(v) for v in value)
                    elif isinstance(value, datetime):
                        value = value.strftime('%Y-%m-%d %H:%M:%S')
                    elif value is None:
                        value = ""
                    row.append(str(value))
                table_data.append(row)
            
            # 创建表格样式
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
        
        # 构建PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _export_docx(self, data: List[Dict], **kwargs) -> bytes:
        """导出为Word文档格式"""
        doc = DocxDocument()
        
        # 标题
        title = doc.add_heading('数据导出报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 导出信息
        doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"总数量: {len(data)}")
        doc.add_paragraph()  # 空行
        
        if data:
            # 创建表格
            headers = list(data[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            # 设置表头
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = str(header)
            
            # 添加数据行
            for item in data:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    value = item.get(header, "")
                    if isinstance(value, (list, tuple)):
                        value = ", ".join(str(v) for v in value)
                    elif isinstance(value, datetime):
                        value = value.strftime('%Y-%m-%d %H:%M:%S')
                    elif value is None:
                        value = ""
                    row_cells[i].text = str(value)
        
        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _export_zip(self, data: List[Dict], **kwargs) -> bytes:
        """导出为ZIP格式（包含多种格式）"""
        buffer = io.BytesIO()
        
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # 添加JSON格式
            json_data = self._export_json(data)
            zip_file.writestr('export.json', json_data)
            
            # 添加CSV格式
            csv_data = self._export_csv(data)
            zip_file.writestr('export.csv', csv_data)
            
            # 添加XML格式
            xml_data = self._export_xml(data)
            zip_file.writestr('export.xml', xml_data)
            
            # 添加PDF格式
            pdf_data = self._export_pdf(data)
            zip_file.writestr('export.pdf', pdf_data)
            
            # 添加Word格式
            docx_data = self._export_docx(data)
            zip_file.writestr('export.docx', docx_data)
            
            # 添加导出信息
            info = {
                "export_time": datetime.now().isoformat(),
                "total_count": len(data),
                "formats": ["json", "csv", "xml", "pdf", "docx"]
            }
            zip_file.writestr('export_info.json', json.dumps(info, ensure_ascii=False, indent=2))
        
        buffer.seek(0)
        return buffer.getvalue()
    
    def get_export_filename(self, format: str, prefix: str = "export") -> str:
        """生成导出文件名"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{prefix}_{timestamp}.{format}"
    
    def validate_export_format(self, format: str) -> bool:
        """验证导出格式是否支持"""
        return format.lower() in self.supported_formats

# 全局导出服务实例
export_service = ExportService()
